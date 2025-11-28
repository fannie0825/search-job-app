import warnings
import os
warnings.filterwarnings('ignore')
os.environ['STREAMLIT_LOG_LEVEL'] = 'error'

import streamlit as st
import requests
import numpy as np
import pandas as pd
from sklearn.metrics.pairwise import cosine_similarity
import time
from datetime import datetime
import json
import re
from io import BytesIO
import PyPDF2
from docx import Document

st.set_page_config(
    page_title="CareerLens - Executive Dashboard",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    /* Dynamic Color Palette - CSS Variables */
    :root {
        /* Light Mode Colors (Elegant/Minimal) */
        --primary-accent: #00796B;
        --action-accent: #FFC107;
        --bg-main: #FAFAFA;
        --bg-container: #F0F0F0;
        --text-primary: #333333;
        --text-secondary: #666666;
        --success-green: #4CAF50;
    }
    
    /* Dark Mode Colors (Professional/Youthful) */
    [data-theme="dark"] {
        --primary-accent: #4DD0E1;
        --action-accent: #FF9800;
        --bg-main: #1E2124;
        --bg-container: #282C31;
        --text-primary: #E0E0E0;
        --text-secondary: #B0B0B0;
        --success-green: #66BB6A;
    }
    
    /* Override Streamlit dark mode - multiple selectors for compatibility */
    .stApp[data-theme="dark"],
    [data-theme="dark"] .stApp,
    .stApp[data-theme="dark"] :root {
        --primary-accent: #4DD0E1;
        --action-accent: #FF9800;
        --bg-main: #1E2124;
        --bg-container: #282C31;
        --text-primary: #E0E0E0;
        --text-secondary: #B0B0B0;
        --success-green: #66BB6A;
    }
    
    /* Ensure CSS variables work in all contexts */
    html[data-theme="dark"],
    html[data-theme="dark"] :root {
        --primary-accent: #4DD0E1;
        --action-accent: #FF9800;
        --bg-main: #1E2124;
        --bg-container: #282C31;
        --text-primary: #E0E0E0;
        --text-secondary: #B0B0B0;
        --success-green: #66BB6A;
    }
    
    /* Main Header */
    .main-header {
        font-size: 3rem;
        font-weight: bold;
        color: var(--primary-accent);
        text-align: center;
        margin-bottom: 1rem;
        letter-spacing: -0.02em;
    }
    
    .sub-header {
        font-size: 1.2rem;
        text-align: center;
        color: var(--text-secondary);
        margin-bottom: 2rem;
    }
    
    /* Step Hierarchy - Visual Markers */
    .step-marker {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 2.5rem;
        height: 2.5rem;
        border-radius: 50%;
        background-color: var(--primary-accent);
        color: white;
        font-size: 1.2rem;
        font-weight: bold;
        margin-right: 1rem;
        flex-shrink: 0;
    }
    
    .step-container {
        margin: 2rem 0;
        padding: 2rem;
        background-color: var(--bg-container);
        border-radius: 12px;
        border: none;
    }
    
    /* Job Card - Minimalist Design */
    .job-card {
        background-color: var(--bg-container);
        padding: 1.5rem;
        border-radius: 12px;
        margin-bottom: 1.5rem;
        border: none;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    
    .job-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
    }
    
    /* Match Score - Action Accent */
    .match-score {
        background-color: var(--action-accent);
        color: white;
        padding: 0.4rem 1rem;
        border-radius: 20px;
        font-weight: bold;
        display: inline-block;
        font-size: 0.9rem;
    }
    
    /* Tags */
    .tag {
        display: inline-block;
        background-color: var(--bg-container);
        color: var(--text-primary);
        padding: 0.3rem 0.8rem;
        border-radius: 12px;
        margin: 0.2rem;
        font-size: 0.85rem;
        border: none;
    }
    
    /* Borderless Input Fields */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stSelectbox > div > div > select {
        border: none !important;
        border-bottom: 1px solid rgba(0, 0, 0, 0.1) !important;
        border-radius: 0 !important;
        background-color: transparent !important;
        padding: 0.5rem 0 !important;
        transition: border-color 0.3s ease !important;
    }
    
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-bottom: 2px solid var(--primary-accent) !important;
        outline: none !important;
        box-shadow: none !important;
    }
    
    /* Dark mode input styling */
    [data-theme="dark"] .stTextInput > div > div > input,
    [data-theme="dark"] .stTextArea > div > div > textarea,
    [data-theme="dark"] .stSelectbox > div > div > select {
        border-bottom: 1px solid rgba(255, 255, 255, 0.2) !important;
        color: var(--text-primary) !important;
    }
    
    [data-theme="dark"] .stTextInput > div > div > input:focus,
    [data-theme="dark"] .stTextArea > div > div > textarea:focus {
        border-bottom: 2px solid var(--primary-accent) !important;
    }
    
    /* Buttons - Rounded with Primary Accent */
    .stButton > button {
        border-radius: 8px !important;
        border: none !important;
        font-weight: 500 !important;
        transition: all 0.2s ease !important;
    }
    
    .stButton > button[kind="primary"] {
        background-color: var(--primary-accent) !important;
        color: white !important;
    }
    
    .stButton > button[kind="primary"]:hover {
        background-color: var(--primary-accent) !important;
        opacity: 0.9 !important;
        transform: translateY(-1px) !important;
    }
    
    /* Sliders - Elegant Design */
    /* Slider track background - subtle gray */
    .stSlider [data-baseweb="slider"] > div:first-child {
        background-color: rgba(0, 0, 0, 0.08) !important;
    }
    
    [data-theme="dark"] .stSlider [data-baseweb="slider"] > div:first-child {
        background-color: rgba(255, 255, 255, 0.15) !important;
    }
    
    /* Slider track fill - primary accent color */
    .stSlider [data-baseweb="slider"] > div > div {
        background-color: var(--primary-accent) !important;
    }
    
    /* Target slider track background */
    .stSlider [data-baseweb="slider-track"] {
        background-color: rgba(0, 0, 0, 0.08) !important;
    }
    
    [data-theme="dark"] .stSlider [data-baseweb="slider-track"] {
        background-color: rgba(255, 255, 255, 0.15) !important;
    }
    
    /* Slider thumb/handle - primary accent with hover effect */
    .stSlider [data-baseweb="slider-thumb"],
    .stSlider [role="slider"] {
        background-color: var(--primary-accent) !important;
        border-color: var(--primary-accent) !important;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1) !important;
        transition: transform 0.2s ease, box-shadow 0.2s ease !important;
    }
    
    .stSlider [data-baseweb="slider-thumb"]:hover,
    .stSlider [role="slider"]:hover {
        transform: scale(1.1) !important;
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.15) !important;
    }
    
    /* Slider value display */
    .stSlider label {
        color: var(--text-primary) !important;
    }
    
    /* Status Indicators - Profile Ready */
    .profile-ready {
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        color: var(--success-green);
        font-weight: 500;
    }
    
    .profile-ready::before {
        content: "✓";
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 1.2rem;
        height: 1.2rem;
        border-radius: 50%;
        background-color: var(--success-green);
        color: white;
        font-size: 0.8rem;
        font-weight: bold;
    }
    
    /* Match Score Display - Prominent with Action Accent */
    .match-score-display {
        font-size: 2rem;
        font-weight: bold;
        color: var(--action-accent);
        text-align: center;
    }
    
    /* Step 3 Containers - Automatic/Manual Search */
    .step3-container {
        background-color: var(--bg-container);
        padding: 1.5rem;
        border-radius: 12px;
        margin: 1rem 0;
    }
    
    /* Alternative: Style columns directly for Step 3 */
    div[data-testid="column"]:has(.step3-container) {
        padding: 0.5rem;
    }
    
    /* Ensure container background shows through */
    .step3-container > * {
        position: relative;
        z-index: 1;
    }
    
    /* Minimal Dividers */
    hr {
        border: none;
        border-top: 1px solid rgba(0, 0, 0, 0.08);
        margin: 2rem 0;
    }
    
    [data-theme="dark"] hr {
        border-top: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    /* Remove default Streamlit borders and shadows */
    .element-container {
        border: none !important;
        box-shadow: none !important;
    }
    
    /* Card-like containers for Step 3 */
    .matching-container {
        background-color: var(--bg-container);
        padding: 1.5rem;
        border-radius: 12px;
        margin-bottom: 1rem;
    }
    
    /* Ensure proper spacing and layout */
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    
    /* Improve form styling */
    .stForm {
        border: none;
        background-color: transparent;
    }
    
    /* File uploader styling */
    .stFileUploader > div {
        border: 1px dashed rgba(0, 0, 0, 0.2);
        border-radius: 8px;
        padding: 1rem;
        background-color: var(--bg-container);
    }
    
    [data-theme="dark"] .stFileUploader > div {
        border-color: rgba(255, 255, 255, 0.2);
    }
    
    /* Expander styling */
    .streamlit-expanderHeader {
        color: var(--text-primary);
        font-weight: 500;
    }
    
    /* Metric styling */
    [data-testid="stMetricValue"] {
        color: var(--text-primary);
    }
    
    [data-testid="stMetricLabel"] {
        color: var(--text-secondary);
    }
    
    /* Info/Warning/Success boxes */
    .stAlert {
        border-radius: 8px;
        border: none;
    }
    
    /* Number input styling */
    .stNumberInput > div > div > input {
        border: none !important;
        border-bottom: 1px solid rgba(0, 0, 0, 0.1) !important;
        border-radius: 0 !important;
        background-color: transparent !important;
    }
    
    [data-theme="dark"] .stNumberInput > div > div > input {
        border-bottom: 1px solid rgba(255, 255, 255, 0.2) !important;
        color: var(--text-primary) !important;
    }
    
    .stNumberInput > div > div > input:focus {
        border-bottom: 2px solid var(--primary-accent) !important;
        outline: none !important;
    }
    
    /* Executive Dashboard Styles */
    .dashboard-header {
        font-size: 2.5rem;
        font-weight: 700;
        color: var(--primary-accent);
        margin-bottom: 0.5rem;
        letter-spacing: -0.02em;
    }
    
    .dashboard-subtitle {
        font-size: 1rem;
        color: var(--text-secondary);
        margin-bottom: 2rem;
    }
    
    .metric-card {
        background-color: var(--bg-container);
        padding: 1.5rem;
        border-radius: 12px;
        border: 1px solid rgba(0, 0, 0, 0.05);
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    
    .metric-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
    }
    
    .metric-value {
        font-size: 2rem;
        font-weight: bold;
        color: var(--primary-accent);
        margin: 0.5rem 0;
    }
    
    .metric-label {
        font-size: 0.9rem;
        color: var(--text-secondary);
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }
    
    .job-table-row {
        cursor: pointer;
        transition: background-color 0.2s ease;
        border-bottom: 1px solid rgba(0, 0, 0, 0.05);
    }
    
    .job-table-row:hover {
        background-color: var(--bg-container);
    }
    
    .match-score-badge {
        display: inline-flex;
        align-items: center;
        padding: 0.4rem 0.8rem;
        border-radius: 20px;
        font-weight: 600;
        font-size: 0.9rem;
    }
    
    .match-score-high {
        background-color: #4CAF50;
        color: white;
    }
    
    .match-score-medium {
        background-color: #FFC107;
        color: #333;
    }
    
    .match-score-low {
        background-color: #FF9800;
        color: white;
    }
    
    .expandable-details {
        background-color: var(--bg-container);
        padding: 1.5rem;
        border-radius: 8px;
        margin-top: 1rem;
        border-left: 4px solid var(--primary-accent);
    }
    
    .match-breakdown {
        display: flex;
        gap: 2rem;
        margin: 1rem 0;
    }
    
    .match-type {
        flex: 1;
        padding: 1rem;
        background-color: var(--bg-main);
        border-radius: 8px;
    }
    
    .match-type-label {
        font-size: 0.85rem;
        color: var(--text-secondary);
        text-transform: uppercase;
        margin-bottom: 0.5rem;
    }
    
    .match-type-value {
        font-size: 1.5rem;
        font-weight: bold;
        color: var(--primary-accent);
    }
    
    /* Sidebar styling */
    .sidebar-section {
        margin-bottom: 2rem;
        padding-bottom: 1.5rem;
        border-bottom: 1px solid rgba(0, 0, 0, 0.1);
    }
    
    [data-theme="dark"] .sidebar-section {
        border-bottom: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    /* Table styling */
    .dataframe {
        width: 100%;
        border-collapse: collapse;
    }
    
    .dataframe th {
        background-color: var(--bg-container);
        color: var(--text-primary);
        font-weight: 600;
        padding: 1rem;
        text-align: left;
        border-bottom: 2px solid var(--primary-accent);
    }
    
    .dataframe td {
        padding: 1rem;
        border-bottom: 1px solid rgba(0, 0, 0, 0.05);
    }
    
    [data-theme="dark"] .dataframe td {
        border-bottom: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    /* Missing Critical Skill column styling */
    .dataframe td:has-text("Missing Critical Skill"),
    .dataframe [data-column="Missing Critical Skill"] {
        color: #FF6B6B !important;
        font-weight: 500;
    }
    
    [data-theme="dark"] .dataframe td:has-text("Missing Critical Skill"),
    [data-theme="dark"] .dataframe [data-column="Missing Critical Skill"] {
        color: #FF8C8C !important;
    }
</style>
""", unsafe_allow_html=True)

if 'search_history' not in st.session_state:
    st.session_state.search_history = []
if 'jobs_cache' not in st.session_state:
    st.session_state.jobs_cache = {}
if 'embedding_gen' not in st.session_state:
    st.session_state.embedding_gen = None
if 'user_profile' not in st.session_state:
    st.session_state.user_profile = {}
if 'generated_resume' not in st.session_state:
    st.session_state.generated_resume = None
if 'text_gen' not in st.session_state:
    st.session_state.text_gen = None
if 'selected_job' not in st.session_state:
    st.session_state.selected_job = None
if 'show_resume_generator' not in st.session_state:
    st.session_state.show_resume_generator = False
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = None
if 'matched_jobs' not in st.session_state:
    st.session_state.matched_jobs = []
if 'match_score' not in st.session_state:
    st.session_state.match_score = None
if 'missing_keywords' not in st.session_state:
    st.session_state.missing_keywords = None
if 'show_profile_editor' not in st.session_state:
    st.session_state.show_profile_editor = False
if 'use_auto_match' not in st.session_state:
    st.session_state.use_auto_match = False
if 'expanded_job_index' not in st.session_state:
    st.session_state.expanded_job_index = None
if 'industry_filter' not in st.session_state:
    st.session_state.industry_filter = None
if 'salary_min' not in st.session_state:
    st.session_state.salary_min = None
if 'salary_max' not in st.session_state:
    st.session_state.salary_max = None
if 'selected_job_index' not in st.session_state:
    st.session_state.selected_job_index = None
if 'dashboard_ready' not in st.session_state:
    st.session_state.dashboard_ready = False

class APIMEmbeddingGenerator:
    def __init__(self, api_key, endpoint):
        self.api_key = api_key
        # Normalize endpoint: remove trailing slash
        endpoint = endpoint.rstrip('/')
        # Remove /openai if it's already in the endpoint (to avoid duplication)
        if endpoint.endswith('/openai'):
            endpoint = endpoint[:-7]  # Remove '/openai'
        self.endpoint = endpoint
        self.deployment = "text-embedding-3-small"
        self.api_version = "2024-02-01"
        self.url = f"{self.endpoint}/openai/deployments/{self.deployment}/embeddings?api-version={self.api_version}"
        self.headers = {"api-key": self.api_key, "Content-Type": "application/json"}
    
    def get_embedding(self, text):
        try:
            payload = {"input": text, "model": self.deployment}
            response = requests.post(self.url, headers=self.headers, json=payload, timeout=30)
            if response.status_code == 200:
                return response.json()['data'][0]['embedding']
            return None
        except Exception as e:
            st.error(f"Error: {e}")
            return None
    
    def get_embeddings_batch(self, texts, batch_size=10):
        embeddings = []
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            progress = (i + len(batch)) / len(texts)
            progress_bar.progress(progress)
            status_text.text(f"🔄 Generating embeddings: {i + len(batch)}/{len(texts)}")
            
            try:
                payload = {"input": batch, "model": self.deployment}
                response = requests.post(self.url, headers=self.headers, json=payload, timeout=30)
                
                if response.status_code == 200:
                    data = response.json()
                    sorted_data = sorted(data['data'], key=lambda x: x['index'])
                    embeddings.extend([item['embedding'] for item in sorted_data])
            except:
                for text in batch:
                    emb = self.get_embedding(text)
                    if emb:
                        embeddings.append(emb)
        
        progress_bar.empty()
        status_text.empty()
        return embeddings

class AzureOpenAITextGenerator:
    def __init__(self, api_key, endpoint):
        self.api_key = api_key
        # Normalize endpoint: remove trailing slash
        endpoint = endpoint.rstrip('/')
        # Remove /openai if it's already in the endpoint (to avoid duplication)
        if endpoint.endswith('/openai'):
            endpoint = endpoint[:-7]  # Remove '/openai'
        self.endpoint = endpoint
        self.deployment = "gpt-4o-mini"  # or your deployment name
        self.api_version = "2024-02-01"
        self.url = f"{self.endpoint}/openai/deployments/{self.deployment}/chat/completions?api-version={self.api_version}"
        self.headers = {"api-key": self.api_key, "Content-Type": "application/json"}
    
    def generate_resume(self, user_profile, job_posting, raw_resume_text=None):
        """Generate a tailored resume based on user profile and job posting using Context Sandwich approach.
        Returns structured JSON data instead of formatted text."""
        
        # System Instructions
        system_instructions = """You are an expert resume writer with expertise in ATS optimization and career coaching.
Your task is to create a tailored resume by analyzing the job description and adapting the user's profile.
Return ONLY valid JSON - no markdown, no additional text, no code blocks."""

        # Job Description Section
        job_description = f"""JOB POSTING TO MATCH:
Title: {job_posting.get('title', 'N/A')}
Company: {job_posting.get('company', 'N/A')}
Description: {job_posting.get('description', 'N/A')}
Required Skills: {', '.join(job_posting.get('skills', []))}"""

        # Structured Profile Section
        structured_profile = f"""STRUCTURED PROFILE:
Name: {user_profile.get('name', 'N/A')}
Email: {user_profile.get('email', 'N/A')}
Phone: {user_profile.get('phone', 'N/A')}
Location: {user_profile.get('location', 'N/A')}
LinkedIn: {user_profile.get('linkedin', 'N/A')}
Portfolio: {user_profile.get('portfolio', 'N/A')}
Summary: {user_profile.get('summary', 'N/A')}
Experience: {user_profile.get('experience', 'N/A')}
Education: {user_profile.get('education', 'N/A')}
Skills: {user_profile.get('skills', 'N/A')}
Certifications: {user_profile.get('certifications', 'N/A')}"""

        # Raw Original Resume (if available)
        raw_resume_section = ""
        if raw_resume_text:
            raw_resume_section = f"\n\nORIGINAL RESUME TEXT (for reference and context):\n{raw_resume_text[:3000]}"  # Limit to avoid token limits

        # Context Sandwich: System Instructions + Job Description + (Structured Profile + Raw Resume)
        prompt = f"""{system_instructions}

{job_description}

{structured_profile}{raw_resume_section}

INSTRUCTIONS:
1. Analyze the job posting requirements and identify key skills, technologies, and qualifications needed
2. Tailor the user's profile to match the job description by:
   - Rewriting the summary to emphasize relevant experience
   - Highlighting skills that match the job requirements
   - Rewriting experience bullet points to emphasize relevant achievements
   - Using keywords from the job description for ATS optimization
3. Focus on achievements and measurable results
4. Maintain accuracy - only use information from the provided profile

Return your response as a JSON object with this exact structure:
{{
  "header": {{
    "name": "Full Name",
    "title": "Professional Title (tailored to job)",
    "email": "email@example.com",
    "phone": "phone number",
    "location": "City, State/Country",
    "linkedin": "LinkedIn URL or empty string",
    "portfolio": "Portfolio URL or empty string"
  }},
  "summary": "2-3 sentence professional summary tailored to the job description, emphasizing relevant experience and skills",
  "skills_highlighted": ["Skill 1", "Skill 2", "Skill 3", ...],
  "experience": [
    {{
      "company": "Company Name",
      "title": "Job Title",
      "dates": "Date Range",
      "bullets": [
        "Rewritten bullet point emphasizing relevant achievement...",
        "Another tailored bullet point..."
      ]
    }}
  ],
  "education": "Education details formatted as text",
  "certifications": "Certifications, awards, or other achievements formatted as text"
}}

IMPORTANT: Return ONLY the JSON object, no markdown code blocks, no additional text."""
        
        try:
            payload = {
                "messages": [
                    {"role": "system", "content": system_instructions},
                    {"role": "user", "content": prompt}
                ],
                "max_tokens": 3000,
                "temperature": 0.7,
                "response_format": {"type": "json_object"}  # Force JSON output
            }
            
            response = requests.post(self.url, headers=self.headers, json=payload, timeout=60)
            
            if response.status_code == 200:
                result = response.json()
                content = result['choices'][0]['message']['content']
                
                # Parse JSON response
                try:
                    # Remove markdown code blocks if present
                    content = content.strip()
                    if content.startswith("```"):
                        lines = content.split('\n')
                        content = '\n'.join(lines[1:-1]) if lines[-1].startswith('```') else '\n'.join(lines[1:])
                    
                    resume_data = json.loads(content)
                    return resume_data
                except json.JSONDecodeError as e:
                    # Try to extract JSON from response
                    json_match = re.search(r'\{.*\}', content, re.DOTALL)
                    if json_match:
                        resume_data = json.loads(json_match.group())
                        return resume_data
                    else:
                        st.error(f"Could not parse JSON response: {e}")
                        return None
            else:
                st.error(f"API Error: {response.status_code} - {response.text}")
                return None
        except Exception as e:
            st.error(f"Error generating resume: {e}")
            return None
    
    def calculate_match_score(self, resume_content, job_description, embedding_generator):
        """Calculate match score between resume and job description, and identify missing keywords"""
        try:
            # Create embeddings for resume and job description
            resume_embedding = embedding_generator.get_embedding(resume_content)
            job_embedding = embedding_generator.get_embedding(job_description)
            
            if not resume_embedding or not job_embedding:
                return None, None
            
            # Calculate cosine similarity
            resume_emb = np.array(resume_embedding).reshape(1, -1)
            job_emb = np.array(job_embedding).reshape(1, -1)
            similarity = cosine_similarity(resume_emb, job_emb)[0][0]
            match_score = float(similarity)
            
            # Extract keywords from job description using AI
            # Use up to 8000 characters for keyword extraction (enough for most JDs while staying within API limits)
            job_desc_for_keywords = job_description[:8000] if len(job_description) > 8000 else job_description
            if len(job_description) > 8000:
                job_desc_for_keywords += "\n\n[Description truncated for keyword extraction - full description available for matching]"
            
            keyword_prompt = f"""Extract the most important technical skills, tools, technologies, and qualifications mentioned in this job description. 
Return ONLY a JSON object with a "keywords" array, no additional text.

Job Description:
{job_desc_for_keywords}

Return format: {{"keywords": ["keyword1", "keyword2", "keyword3", ...]}}"""
            
            payload = {
                "messages": [
                    {"role": "system", "content": "You are a keyword extraction expert. Extract only the most important technical and professional keywords. Return JSON with a 'keywords' array."},
                    {"role": "user", "content": keyword_prompt}
                ],
                "max_tokens": 500,
                "temperature": 0.3,
                "response_format": {"type": "json_object"}
            }
            
            response = requests.post(self.url, headers=self.headers, json=payload, timeout=30)
            
            missing_keywords = []
            if response.status_code == 200:
                try:
                    result = response.json()
                    content = result['choices'][0]['message']['content']
                    # Try to parse keywords
                    keyword_data = json.loads(content)
                    job_keywords = keyword_data.get('keywords', [])
                    
                    # Check which keywords are missing from resume
                    resume_lower = resume_content.lower()
                    for keyword in job_keywords:
                        if isinstance(keyword, str) and keyword.lower() not in resume_lower:
                            missing_keywords.append(keyword)
                except Exception as e:
                    # If keyword extraction fails, continue without missing keywords
                    pass
            
            return match_score, missing_keywords[:10]  # Limit to top 10 missing keywords
            
        except Exception as e:
            st.warning(f"Could not calculate match score: {e}")
            return None, None

class IndeedScraperAPI:
    def __init__(self, api_key):
        self.api_key = api_key
        self.url = "https://indeed-scraper-api.p.rapidapi.com/api/job"
        self.headers = {
            'Content-Type': 'application/json',
            'x-rapidapi-host': 'indeed-scraper-api.p.rapidapi.com',
            'x-rapidapi-key': api_key
        }
    
    def search_jobs(self, query, location="Hong Kong", max_rows=15, job_type="fulltime", country="hk"):
        payload = {
            "scraper": {
                "maxRows": max_rows,
                "query": query,
                "location": location,
                "jobType": job_type,
                "radius": "50",
                "sort": "relevance",
                "fromDays": "7",
                "country": country
            }
        }
        
        try:
            response = requests.post(self.url, headers=self.headers, json=payload, timeout=60)
            
            if response.status_code == 201:
                data = response.json()
                jobs = []
                
                if 'returnvalue' in data and 'data' in data['returnvalue']:
                    job_list = data['returnvalue']['data']
                    
                    for job_data in job_list:
                        parsed_job = self._parse_job(job_data)
                        if parsed_job:
                            jobs.append(parsed_job)
                
                return jobs
            else:
                st.error(f"API Error: {response.status_code}")
                return []
                
        except Exception as e:
            st.error(f"Error: {e}")
            return []
    
    def _parse_job(self, job_data):
        try:
            location_data = job_data.get('location', {})
            location = location_data.get('formattedAddressShort') or location_data.get('city', 'Hong Kong')
            
            job_types = job_data.get('jobType', [])
            job_type = ', '.join(job_types) if job_types else 'Full-time'
            
            benefits = job_data.get('benefits', [])
            attributes = job_data.get('attributes', [])
            
            # Get full description without truncation
            full_description = job_data.get('descriptionText', 'No description')
            # Store full description, but limit to 50000 chars to prevent memory issues
            description = full_description[:50000] if len(full_description) > 50000 else full_description
            
            return {
                'title': job_data.get('title', 'N/A'),
                'company': job_data.get('companyName', 'N/A'),
                'location': location,
                'description': description,
                'salary': 'Not specified',
                'job_type': job_type,
                'url': job_data.get('jobUrl', '#'),
                'posted_date': job_data.get('age', 'Recently'),
                'benefits': benefits[:5],
                'skills': attributes[:10],
                'company_rating': job_data.get('rating', {}).get('rating', 0),
                'is_remote': job_data.get('isRemote', False)
            }
        except:
            return None

class SemanticJobSearch:
    def __init__(self, embedding_generator):
        self.embedding_gen = embedding_generator
        self.job_embeddings = []
        self.jobs = []
    
    def index_jobs(self, jobs):
        self.jobs = jobs
        job_texts = [
            f"{job['title']} at {job['company']}. {job['description']} Skills: {', '.join(job['skills'][:5])}"
            for job in jobs
        ]
        
        st.info(f"📊 Indexing {len(jobs)} jobs...")
        self.job_embeddings = self.embedding_gen.get_embeddings_batch(job_texts)
        st.success(f"✅ Indexed {len(self.job_embeddings)} jobs")
    
    def search(self, query, top_k=10):
        if not self.job_embeddings:
            return []
        
        query_embedding = self.embedding_gen.get_embedding(query)
        if not query_embedding:
            return []
        
        query_emb = np.array(query_embedding).reshape(1, -1)
        job_embs = np.array(self.job_embeddings)
        
        similarities = cosine_similarity(query_emb, job_embs)[0]
        top_indices = np.argsort(similarities)[::-1][:top_k]
        
        results = []
        for idx in top_indices:
            results.append({
                'job': self.jobs[idx],
                'similarity_score': float(similarities[idx]),
                'rank': len(results) + 1
            })
        
        return results
    
    def calculate_skill_match(self, user_skills, job_skills):
        """Calculate skill-based match score"""
        if not user_skills or not job_skills:
            return 0.0, []
        
        # Normalize skills to lowercase for comparison
        user_skills_lower = [s.lower().strip() for s in str(user_skills).split(',') if s.strip()]
        job_skills_lower = [s.lower().strip() for s in job_skills if isinstance(s, str) and s.strip()]
        
        if not user_skills_lower or not job_skills_lower:
            return 0.0, []
        
        # Find matching skills
        matched_skills = []
        for job_skill in job_skills_lower:
            for user_skill in user_skills_lower:
                if job_skill in user_skill or user_skill in job_skill:
                    matched_skills.append(job_skill)
                    break
        
        # Calculate match percentage
        match_score = len(matched_skills) / len(job_skills_lower) if job_skills_lower else 0.0
        missing_skills = [s for s in job_skills_lower if s not in matched_skills]
        
        return min(match_score, 1.0), missing_skills[:5]  # Cap at 1.0 and limit missing skills

def get_embedding_generator():
    if st.session_state.embedding_gen is None:
        # Use secrets instead of hardcoded values
        AZURE_OPENAI_API_KEY = st.secrets["AZURE_OPENAI_API_KEY"]
        AZURE_OPENAI_ENDPOINT = st.secrets["AZURE_OPENAI_ENDPOINT"]
        st.session_state.embedding_gen = APIMEmbeddingGenerator(AZURE_OPENAI_API_KEY, AZURE_OPENAI_ENDPOINT)
    return st.session_state.embedding_gen

def get_job_scraper():
    # Use secrets instead of hardcoded values
    RAPIDAPI_KEY = st.secrets["RAPIDAPI_KEY"]
    return IndeedScraperAPI(RAPIDAPI_KEY)

def get_text_generator():
    if st.session_state.text_gen is None:
        AZURE_OPENAI_API_KEY = st.secrets["AZURE_OPENAI_API_KEY"]
        AZURE_OPENAI_ENDPOINT = st.secrets["AZURE_OPENAI_ENDPOINT"]
        st.session_state.text_gen = AzureOpenAITextGenerator(AZURE_OPENAI_API_KEY, AZURE_OPENAI_ENDPOINT)
    return st.session_state.text_gen

def display_job_card(result, index):
    job = result['job']
    score = result['similarity_score']
    
    remote_badge = "🏠 Remote" if job['is_remote'] else ""
    rating = job['company_rating']
    stars = "⭐" * int(rating) if rating > 0 else ""
    
    st.markdown(f"""
    <div class="job-card">
        <div style="display: flex; justify-content: space-between; align-items: start; margin-bottom: 1rem;">
            <div style="flex-grow: 1;">
                <h3 style="margin: 0; color: var(--primary-accent);">#{index} {job['title']}</h3>
                <p style="margin: 0.5rem 0; color: var(--text-secondary); font-size: 0.95rem;">
                    🏢 <strong>{job['company']}</strong> {stars} • 📍 {job['location']} {remote_badge}
                </p>
            </div>
            <div class="match-score">
                {score:.1%} Match
            </div>
        </div>
        <div style="display: flex; gap: 1rem; flex-wrap: wrap; margin-bottom: 0.5rem; color: var(--text-secondary);">
            <span>⏰ {job['job_type']}</span>
            <span>💰 {job['salary']}</span>
            <span>📅 {job['posted_date']}</span>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        if job['benefits']:
            st.write("**Benefits:**")
            for benefit in job['benefits']:
                st.markdown(f'<span class="tag">✓ {benefit}</span>', unsafe_allow_html=True)
    
    with col2:
        if job['skills']:
            st.write("**Skills:**")
            skills_text = " ".join([f'<span class="tag">{skill}</span>' for skill in job['skills'][:8]])
            st.markdown(skills_text, unsafe_allow_html=True)
    
    st.write("")
    
    col1, col2 = st.columns([4, 1])
    
    with col1:
        with st.expander("📝 View Full Description", expanded=False):
            # Display full description with proper formatting
            description_text = job['description']
            if len(description_text) > 10000:
                st.info(f"📄 Full description ({len(description_text):,} characters)")
                # Use text area for very long descriptions to allow scrolling
                st.text_area(
                    "Job Description",
                    value=description_text,
                    height=400,
                    key=f"desc_{index}",
                    label_visibility="collapsed"
                )
            else:
                st.write(description_text)
    
    with col2:
        col2a, col2b = st.columns(2)
        with col2a:
            if job['url'] != '#':
                st.link_button("Apply →", job['url'], use_container_width=True)
        with col2b:
            if st.button("📄 Resume", key=f"resume_{index}", use_container_width=True, type="primary"):
                st.session_state.selected_job = job
                st.session_state.show_resume_generator = True
                st.rerun()

def extract_text_from_resume(uploaded_file):
    """Extract text from uploaded resume file (PDF, DOCX, or TXT)"""
    try:
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        if file_type == 'pdf':
            # Extract text from PDF
            uploaded_file.seek(0)  # Reset file pointer
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        
        elif file_type == 'docx':
            # Extract text from DOCX
            uploaded_file.seek(0)  # Reset file pointer
            doc = Document(uploaded_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        
        elif file_type == 'txt':
            # Read text file
            uploaded_file.seek(0)  # Reset file pointer
            text = str(uploaded_file.read(), "utf-8")
            return text
        
        else:
            st.error(f"Unsupported file type: {file_type}. Please upload PDF, DOCX, or TXT.")
            return None
            
    except Exception as e:
        st.error(f"Error extracting text from resume: {e}")
        return None

def extract_profile_from_resume(resume_text):
    """Use Azure OpenAI to extract structured profile information from resume text"""
    try:
        text_gen = get_text_generator()
        
        prompt = f"""You are an expert at parsing resumes. Extract structured information from the following resume text.

RESUME TEXT:
{resume_text}

Please extract and return the following information in JSON format:
{{
    "name": "Full name",
    "email": "Email address",
    "phone": "Phone number",
    "location": "City, State/Country",
    "linkedin": "LinkedIn URL if mentioned",
    "portfolio": "Portfolio/website URL if mentioned",
    "summary": "Professional summary or objective (2-3 sentences)",
    "experience": "Work experience in chronological order with job titles, companies, dates, and key achievements (formatted as bullet points)",
    "education": "Education details including degrees, institutions, and graduation dates",
    "skills": "Comma-separated list of technical and soft skills",
    "certifications": "Professional certifications, awards, publications, or other achievements"
}}

Important:
- If information is not found, use "N/A" or empty string
- Format experience with clear job titles, companies, dates, and bullet points for achievements
- Extract all relevant skills mentioned
- Keep the summary concise but informative
- Return ONLY valid JSON, no additional text or markdown"""
        
        payload = {
            "messages": [
                {"role": "system", "content": "You are a resume parser. Extract structured information and return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 2000,
            "temperature": 0.3
        }
        
        response = requests.post(
            text_gen.url,
            headers=text_gen.headers,
            json=payload,
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content']
            
            # Try to extract JSON from the response
            # Sometimes the model returns JSON wrapped in markdown code blocks
            content = content.strip()
            if content.startswith("```"):
                # Remove markdown code blocks
                lines = content.split('\n')
                content = '\n'.join(lines[1:-1]) if lines[-1].startswith('```') else '\n'.join(lines[1:])
            
            try:
                profile_data = json.loads(content)
                return profile_data
            except json.JSONDecodeError:
                # Try to find JSON in the response
                json_match = re.search(r'\{.*\}', content, re.DOTALL)
                if json_match:
                    profile_data = json.loads(json_match.group())
                    return profile_data
                else:
                    st.error("Could not parse extracted profile data. Please try again.")
                    return None
        else:
            st.error(f"API Error: {response.status_code} - {response.text}")
            return None
            
    except Exception as e:
        st.error(f"Error extracting profile: {e}")
        return None

def display_user_profile():
    """Display and edit user profile"""
    st.header("👤 Your Profile")
    st.caption("Fill in your information to generate tailored resumes")
    
    # Resume upload section
    st.markdown("---")
    st.subheader("📄 Upload Your Resume (Optional)")
    st.caption("Upload your resume to automatically extract your information")
    
    uploaded_file = st.file_uploader(
        "Choose a resume file",
        type=['pdf', 'docx', 'txt'],
        help="Supported formats: PDF, DOCX, TXT"
    )
    
    if uploaded_file is not None:
        if st.button("🔍 Extract Information from Resume", type="primary", use_container_width=True):
            with st.spinner("📖 Reading resume and extracting information..."):
                # Extract text from resume
                resume_text = extract_text_from_resume(uploaded_file)
                
                if resume_text:
                    # Store resume text for job matching
                    st.session_state.resume_text = resume_text
                    st.success(f"✅ Extracted {len(resume_text)} characters from resume")
                    
                    # Show extracted text preview
                    with st.expander("📝 Preview Extracted Text"):
                        st.text(resume_text[:1000] + "..." if len(resume_text) > 1000 else resume_text)
                    
                    # Extract structured information
                    with st.spinner("🤖 Using AI to extract structured information..."):
                        profile_data = extract_profile_from_resume(resume_text)
                        
                        if profile_data:
                            # Update session state with extracted data
                            st.session_state.user_profile = {
                                'name': profile_data.get('name', ''),
                                'email': profile_data.get('email', ''),
                                'phone': profile_data.get('phone', ''),
                                'location': profile_data.get('location', ''),
                                'linkedin': profile_data.get('linkedin', ''),
                                'portfolio': profile_data.get('portfolio', ''),
                                'summary': profile_data.get('summary', ''),
                                'experience': profile_data.get('experience', ''),
                                'education': profile_data.get('education', ''),
                                'skills': profile_data.get('skills', ''),
                                'certifications': profile_data.get('certifications', '')
                            }
                            st.success("✅ Profile information extracted successfully! Review and edit below.")
                            st.balloons()
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.warning("⚠️ Could not extract structured information. Please fill in manually.")
    
    st.markdown("---")
    
    with st.form("user_profile_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            name = st.text_input("Full Name", value=st.session_state.user_profile.get('name', ''))
            email = st.text_input("Email", value=st.session_state.user_profile.get('email', ''))
            phone = st.text_input("Phone", value=st.session_state.user_profile.get('phone', ''))
        
        with col2:
            location = st.text_input("Location", value=st.session_state.user_profile.get('location', ''))
            linkedin = st.text_input("LinkedIn URL", value=st.session_state.user_profile.get('linkedin', ''))
            portfolio = st.text_input("Portfolio/Website", value=st.session_state.user_profile.get('portfolio', ''))
        
        summary = st.text_area(
            "Professional Summary",
            value=st.session_state.user_profile.get('summary', ''),
            height=100,
            placeholder="Brief overview of your professional background..."
        )
        
        experience = st.text_area(
            "Work Experience",
            value=st.session_state.user_profile.get('experience', ''),
            height=150,
            placeholder="List your work experience with job titles, companies, dates, and key achievements..."
        )
        
        education = st.text_area(
            "Education",
            value=st.session_state.user_profile.get('education', ''),
            height=100,
            placeholder="Degrees, institutions, graduation dates..."
        )
        
        skills = st.text_area(
            "Skills",
            value=st.session_state.user_profile.get('skills', ''),
            height=80,
            placeholder="List your technical and soft skills (comma-separated)..."
        )
        
        certifications = st.text_area(
            "Certifications & Awards",
            value=st.session_state.user_profile.get('certifications', ''),
            height=80,
            placeholder="Professional certifications, awards, publications..."
        )
        
        submitted = st.form_submit_button("💾 Save Profile", use_container_width=True, type="primary")
        
        if submitted:
            st.session_state.user_profile = {
                'name': name,
                'email': email,
                'phone': phone,
                'location': location,
                'linkedin': linkedin,
                'portfolio': portfolio,
                'summary': summary,
                'experience': experience,
                'education': education,
                'skills': skills,
                'certifications': certifications
            }
            st.success("✅ Profile saved successfully!")
            time.sleep(1)
            st.rerun()

def render_structured_resume_editor(resume_data):
    """Render structured resume JSON in editable Streamlit form"""
    if not resume_data:
        return None
    
    edited_data = {}
    
    st.subheader("📋 Your Tailored Resume")
    st.caption("Edit the sections below to customize your resume")
    
    # Header Section
    with st.expander("👤 Header Information", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            edited_data['header'] = {
                'name': st.text_input("Full Name", value=resume_data.get('header', {}).get('name', ''), key='resume_name'),
                'title': st.text_input("Professional Title", value=resume_data.get('header', {}).get('title', ''), key='resume_title'),
                'email': st.text_input("Email", value=resume_data.get('header', {}).get('email', ''), key='resume_email'),
                'phone': st.text_input("Phone", value=resume_data.get('header', {}).get('phone', ''), key='resume_phone'),
            }
        with col2:
            edited_data['header']['location'] = st.text_input("Location", value=resume_data.get('header', {}).get('location', ''), key='resume_location')
            edited_data['header']['linkedin'] = st.text_input("LinkedIn URL", value=resume_data.get('header', {}).get('linkedin', ''), key='resume_linkedin')
            edited_data['header']['portfolio'] = st.text_input("Portfolio URL", value=resume_data.get('header', {}).get('portfolio', ''), key='resume_portfolio')
    
    # Summary
    edited_data['summary'] = st.text_area(
        "Professional Summary",
        value=resume_data.get('summary', ''),
        height=100,
        key='resume_summary'
    )
    
    # Skills
    skills_list = resume_data.get('skills_highlighted', [])
    skills_text = ', '.join(skills_list) if skills_list else ''
    skills_input = st.text_area(
        "Highlighted Skills (comma-separated)",
        value=skills_text,
        height=60,
        key='resume_skills',
        help="List skills separated by commas"
    )
    edited_data['skills_highlighted'] = [s.strip() for s in skills_input.split(',') if s.strip()]
    
    # Experience
    st.subheader("💼 Work Experience")
    edited_data['experience'] = []
    
    experience_list = resume_data.get('experience', [])
    for i, exp in enumerate(experience_list):
        with st.expander(f"📌 {exp.get('company', 'Company')} - {exp.get('title', 'Position')}", expanded=(i == 0)):
            col1, col2 = st.columns([2, 1])
            with col1:
                company = st.text_input("Company", value=exp.get('company', ''), key=f'exp_company_{i}')
                title = st.text_input("Job Title", value=exp.get('title', ''), key=f'exp_title_{i}')
            with col2:
                dates = st.text_input("Date Range", value=exp.get('dates', ''), key=f'exp_dates_{i}')
            
            st.write("**Key Achievements:**")
            bullets = exp.get('bullets', [])
            edited_bullets = []
            for j, bullet in enumerate(bullets):
                bullet_text = st.text_area(
                    f"Bullet {j+1}",
                    value=bullet,
                    height=60,
                    key=f'exp_bullet_{i}_{j}'
                )
                if bullet_text.strip():
                    edited_bullets.append(bullet_text.strip())
            
            # Allow adding new bullets
            if st.button(f"➕ Add Bullet Point", key=f'add_bullet_{i}'):
                edited_bullets.append("")
                st.rerun()
            
            edited_data['experience'].append({
                'company': company,
                'title': title,
                'dates': dates,
                'bullets': edited_bullets
            })
    
    # Education
    edited_data['education'] = st.text_area(
        "Education",
        value=resume_data.get('education', ''),
        height=100,
        key='resume_education'
    )
    
    # Certifications
    edited_data['certifications'] = st.text_area(
        "Certifications & Awards",
        value=resume_data.get('certifications', ''),
        height=100,
        key='resume_certifications'
    )
    
    return edited_data

def generate_docx_from_json(resume_data, filename="resume.docx"):
    """Generate a professional .docx file from structured resume JSON"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Header Section
        header = resume_data.get('header', {})
        if header.get('name'):
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(header['name'])
            name_run.font.size = Pt(18)
            name_run.font.bold = True
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact information
        contact_info = []
        if header.get('email'):
            contact_info.append(header['email'])
        if header.get('phone'):
            contact_info.append(header['phone'])
        if header.get('location'):
            contact_info.append(header['location'])
        if header.get('linkedin'):
            contact_info.append(header['linkedin'])
        if header.get('portfolio'):
            contact_info.append(header['portfolio'])
        
        if contact_info:
            contact_para = doc.add_paragraph(' | '.join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()  # Spacing
        
        # Professional Title
        if header.get('title'):
            title_para = doc.add_paragraph(header['title'])
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.runs[0].font.size = Pt(12)
            title_para.runs[0].italic = True
            doc.add_paragraph()  # Spacing
        
        # Summary
        if resume_data.get('summary'):
            doc.add_heading('Professional Summary', level=2)
            summary_para = doc.add_paragraph(resume_data['summary'])
            summary_para.runs[0].font.size = Pt(11)
            doc.add_paragraph()  # Spacing
        
        # Skills
        skills = resume_data.get('skills_highlighted', [])
        if skills:
            doc.add_heading('Key Skills', level=2)
            skills_text = ' • '.join(skills)
            skills_para = doc.add_paragraph(skills_text)
            skills_para.runs[0].font.size = Pt(11)
            doc.add_paragraph()  # Spacing
        
        # Experience
        experience = resume_data.get('experience', [])
        if experience:
            doc.add_heading('Professional Experience', level=2)
            for exp in experience:
                # Company and Title
                exp_header = doc.add_paragraph()
                exp_header.add_run(exp.get('title', '')).bold = True
                if exp.get('company'):
                    exp_header.add_run(f" at {exp['company']}")
                if exp.get('dates'):
                    exp_header.add_run(f" | {exp['dates']}").italic = True
                
                # Bullet points
                bullets = exp.get('bullets', [])
                for bullet in bullets:
                    if bullet.strip():
                        bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                        bullet_para.runs[0].font.size = Pt(10)
                
                doc.add_paragraph()  # Spacing between experiences
        
        # Education
        if resume_data.get('education'):
            doc.add_heading('Education', level=2)
            edu_para = doc.add_paragraph(resume_data['education'])
            edu_para.runs[0].font.size = Pt(11)
            doc.add_paragraph()  # Spacing
        
        # Certifications
        if resume_data.get('certifications'):
            doc.add_heading('Certifications & Awards', level=2)
            cert_para = doc.add_paragraph(resume_data['certifications'])
            cert_para.runs[0].font.size = Pt(11)
        
        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
        
    except Exception as e:
        st.error(f"Error generating DOCX: {e}")
        return None

def display_match_score_feedback(match_score, missing_keywords, job_title):
    """Display match score and feedback to user"""
    if match_score is None:
        return
    
    st.markdown("---")
    st.subheader("🎯 Resume Match Analysis")
    
    # Match score with color coding
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        score_percent = match_score * 100
        
        if score_percent >= 80:
            score_color = "🟢"
            feedback = "Excellent match! Your resume aligns well with this position."
        elif score_percent >= 60:
            score_color = "🟡"
            feedback = "Good match. Consider adding more relevant keywords."
        else:
            score_color = "🔴"
            feedback = "Moderate match. Your resume may need more tailoring."
        
        st.markdown(f"""
        <div style="text-align: center; margin: 1rem 0;">
            <div class="match-score-display">{score_percent:.0f}%</div>
            <p style="color: var(--text-secondary); margin-top: 0.5rem;">Match Score</p>
        </div>
        """, unsafe_allow_html=True)
        st.caption(f"**Analysis:** {feedback}")
    
    # Missing keywords
    if missing_keywords:
        st.warning(f"**Missing Keywords:** {', '.join(missing_keywords[:5])}")
        if len(missing_keywords) > 5:
            with st.expander(f"See all {len(missing_keywords)} missing keywords"):
                st.write(', '.join(missing_keywords))
        
        st.info("💡 **Tip:** Consider adding these keywords to your resume if you have experience with them. Be honest - only include skills you actually possess.")
    else:
        st.success("✅ Great! Your resume includes the key keywords from the job description.")
    
    # Feedback
    st.caption(f"**Analysis:** {feedback}")

def display_resume_generator():
    """Display the resume generator interface with structured resume editing"""
    if st.session_state.selected_job is None:
        st.warning("No job selected. Please select a job first.")
        if st.button("← Back to Jobs"):
            st.session_state.show_resume_generator = False
            st.rerun()
        return
    
    job = st.session_state.selected_job
    
    st.markdown('<h1 class="main-header">📄 Resume Generator</h1>', unsafe_allow_html=True)
    
    # Display selected job info
    st.markdown(f"""
    <div class="job-card">
        <h3 style="color: var(--primary-accent); margin: 0;">{job['title']}</h3>
        <p style="margin: 0.5rem 0; color: var(--text-secondary);">🏢 {job['company']} • 📍 {job['location']}</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Check if user profile is complete
    if not st.session_state.user_profile.get('name') or not st.session_state.user_profile.get('experience'):
        st.error("⚠️ Please complete your profile first!")
        if st.button("← Go to Profile"):
            st.session_state.show_resume_generator = False
            st.rerun()
        return
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        st.write("**Your Profile:**", st.session_state.user_profile.get('name', 'N/A'))
    
    with col2:
        if st.button("← Back to Jobs"):
            st.session_state.show_resume_generator = False
            st.session_state.generated_resume = None
            st.session_state.match_score = None
            st.session_state.missing_keywords = None
            st.rerun()
    
    st.markdown("---")
    
    # Generate resume button
    if st.button("🚀 Generate Tailored Resume", type="primary", use_container_width=True):
        with st.spinner("🤖 Creating your personalized resume using AI..."):
            text_gen = get_text_generator()
            # Get raw resume text if available
            raw_resume_text = st.session_state.get('resume_text')
            resume_data = text_gen.generate_resume(
                st.session_state.user_profile, 
                job,
                raw_resume_text=raw_resume_text
            )
            
            if resume_data:
                st.session_state.generated_resume = resume_data
                
                # Calculate match score
                with st.spinner("📊 Analyzing resume match..."):
                    embedding_gen = get_embedding_generator()
                    # Convert resume JSON to text for comparison
                    resume_text = json.dumps(resume_data, indent=2)
                    match_score, missing_keywords = text_gen.calculate_match_score(
                        resume_text,
                        job.get('description', ''),
                        embedding_gen
                    )
                    st.session_state.match_score = match_score
                    st.session_state.missing_keywords = missing_keywords
                
                st.success("✅ Resume generated successfully!")
                st.balloons()
                time.sleep(0.5)
                st.rerun()
            else:
                st.error("❌ Failed to generate resume. Please try again.")
    
    # Display match score if available
    if st.session_state.generated_resume and st.session_state.get('match_score') is not None:
        display_match_score_feedback(
            st.session_state.match_score,
            st.session_state.missing_keywords,
            job['title']
        )
    
    # Display generated resume in structured form
    if st.session_state.generated_resume:
        st.markdown("---")
        
        # Render structured editor
        edited_resume_data = render_structured_resume_editor(st.session_state.generated_resume)
        
        # Update session state with edited data
        if edited_resume_data:
            st.session_state.generated_resume = edited_resume_data
        
        st.markdown("---")
        
        # Download buttons
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            # Download as DOCX
            docx_file = generate_docx_from_json(
                st.session_state.generated_resume,
                filename=f"resume_{job['company']}_{job['title']}.docx"
            )
            if docx_file:
                st.download_button(
                    label="📥 Download as DOCX",
                    data=docx_file,
                    file_name=f"resume_{job['company']}_{job['title']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        
        with col2:
            # Download as JSON
            json_data = json.dumps(st.session_state.generated_resume, indent=2)
            st.download_button(
                label="📥 Download as JSON",
                data=json_data,
                file_name=f"resume_{job['company']}_{job['title']}.json",
                mime="application/json",
                use_container_width=True
            )
        
        with col3:
            # Download as TXT (formatted text version)
            txt_content = format_resume_as_text(st.session_state.generated_resume)
            st.download_button(
                label="📥 Download as TXT",
                data=txt_content,
                file_name=f"resume_{job['company']}_{job['title']}.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        with col4:
            # Apply to job button
            if job['url'] != '#':
                st.link_button(
                    "🚀 Apply to Job",
                    job['url'],
                    use_container_width=True,
                    type="primary"
                )
        
        # Recalculate match score button
        if st.button("🔄 Recalculate Match Score", use_container_width=True):
            with st.spinner("📊 Recalculating match score..."):
                text_gen = get_text_generator()
                embedding_gen = get_embedding_generator()
                resume_text = json.dumps(st.session_state.generated_resume, indent=2)
                match_score, missing_keywords = text_gen.calculate_match_score(
                    resume_text,
                    job.get('description', ''),
                    embedding_gen
                )
                st.session_state.match_score = match_score
                st.session_state.missing_keywords = missing_keywords
                st.rerun()

def render_sidebar():
    """Render CareerLens sidebar with resume upload, market filters, and analyze button"""
    with st.sidebar:
        # Header with icon and title
        st.markdown("""
        <div style="margin-bottom: 2rem;">
            <h2 style="color: #0F62FE; margin-bottom: 0.5rem; display: flex; align-items: center; gap: 0.5rem;">
                🔍 CareerLens
            </h2>
            <p style="color: #666; font-size: 0.9rem; margin: 0;">AI Career Copilot for Hong Kong</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Resume Upload Section
        st.markdown("---")
        st.markdown("### 1. Upload your CV to begin")
        uploaded_file = st.file_uploader(
            "Upload your resume",
            type=['pdf', 'docx'],
            help="We parse your skills and experience to benchmark you against the market.",
            key="careerlens_resume_upload",
            label_visibility="collapsed"
        )
        
        if uploaded_file is not None:
            with st.spinner("📖 Reading resume..."):
                resume_text = extract_text_from_resume(uploaded_file)
                if resume_text:
                    st.session_state.resume_text = resume_text
                    
                    # Extract structured information
                    with st.spinner("🤖 Extracting profile data..."):
                        profile_data = extract_profile_from_resume(resume_text)
                        if profile_data:
                            st.session_state.user_profile = {
                                'name': profile_data.get('name', ''),
                                'email': profile_data.get('email', ''),
                                'phone': profile_data.get('phone', ''),
                                'location': profile_data.get('location', ''),
                                'linkedin': profile_data.get('linkedin', ''),
                                'portfolio': profile_data.get('portfolio', ''),
                                'summary': profile_data.get('summary', ''),
                                'experience': profile_data.get('experience', ''),
                                'education': profile_data.get('education', ''),
                                'skills': profile_data.get('skills', ''),
                                'certifications': profile_data.get('certifications', '')
                            }
                            st.success("✅ Profile extracted!")
        
        # Market Filters Section
        st.markdown("---")
        st.markdown("### 2. Refine Market Scope")
        
        # Target Domains
        target_domains = st.multiselect(
            "Select Target Domains (HK Focus)",
            options=["FinTech", "ESG & Sustainability", "Data Analytics", "Digital Transformation", 
                    "Investment Banking", "Consulting", "Technology", "Healthcare", "Education"],
            default=[],
            key="careerlens_domains"
        )
        st.session_state.target_domains = target_domains
        
        # Salary Expectations
        salary_expectation = st.slider(
            "Min. Monthly Salary Expectation (HKD)",
            min_value=20000,
            max_value=150000,
            value=45000,
            step=5000,
            key="careerlens_salary"
        )
        st.session_state.salary_expectation = salary_expectation
        
        # Primary Action Button
        st.markdown("---")
        analyze_button = st.button(
            "Analyze Profile & Find Matches",
            type="primary",
            use_container_width=True,
            key="careerlens_analyze"
        )
        
        if analyze_button:
            # Fetch jobs and perform matching
            if not st.session_state.resume_text and not st.session_state.user_profile.get('summary'):
                st.error("⚠️ Please upload your CV first!")
            else:
                # Fetch jobs based on filters
                search_query = " ".join(target_domains) if target_domains else "Hong Kong jobs"
                scraper = get_job_scraper()
                
                with st.spinner("🔄 Fetching jobs and analyzing..."):
                    jobs = scraper.search_jobs(search_query, "Hong Kong", 15, "fulltime", "hk")
                    
                    if jobs:
                        st.session_state.jobs_cache = {
                            'jobs': jobs,
                            'count': len(jobs),
                            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            'query': search_query
                        }
                        
                        # Perform semantic matching
                        embedding_gen = get_embedding_generator()
                        search_engine = SemanticJobSearch(embedding_gen)
                        search_engine.index_jobs(jobs)
                        
                        # Build query from resume/profile
                        if st.session_state.resume_text:
                            resume_query = st.session_state.resume_text
                            if st.session_state.user_profile.get('summary'):
                                profile_data = f"{st.session_state.user_profile.get('summary', '')} {st.session_state.user_profile.get('experience', '')} {st.session_state.user_profile.get('skills', '')}"
                                resume_query = f"{resume_query} {profile_data}"
                        else:
                            resume_query = f"{st.session_state.user_profile.get('summary', '')} {st.session_state.user_profile.get('experience', '')} {st.session_state.user_profile.get('skills', '')} {st.session_state.user_profile.get('education', '')}"
                        
                        results = search_engine.search(resume_query, top_k=min(15, len(jobs)))
                        
                        # Calculate skill matches
                        user_skills = st.session_state.user_profile.get('skills', '')
                        for result in results:
                            job_skills = result['job'].get('skills', [])
                            skill_score, missing_skills = search_engine.calculate_skill_match(user_skills, job_skills)
                            result['skill_match_score'] = skill_score
                            result['missing_skills'] = missing_skills
                        
                        st.session_state.matched_jobs = results
                        st.session_state.dashboard_ready = True
                        st.rerun()
                    else:
                        st.error("❌ No jobs found. Please try different filters.")

def display_market_positioning_profile(matched_jobs, user_profile):
    """Display Market Positioning Profile with 4 key metrics"""
    if not matched_jobs:
        return
    
    # Get user name or use placeholder
    user_name = user_profile.get('name', 'Professional')
    if not user_name or user_name == 'N/A':
        user_name = 'Professional'
    
    st.markdown(f"### Welcome, {user_name}. Here is your market positioning snapshot.")
    
    # Calculate metrics
    # Metric 1: Estimated Market Salary Band
    # Simulate salary calculation based on matched jobs
    salary_min = 45000
    salary_max = 55000
    
    # Metric 2: Target Role Seniority
    # Analyze job titles to determine seniority level
    job_titles = [r['job'].get('title', '') for r in matched_jobs[:5]]
    seniority = "Mid-Senior Level"  # Default, could be enhanced with AI analysis
    
    # Metric 3: Top Skill Gap
    user_skills = user_profile.get('skills', '')
    all_job_skills = []
    for result in matched_jobs:
        all_job_skills.extend(result['job'].get('skills', []))
    
    # Find most common missing skill
    user_skills_list = [s.lower().strip() for s in str(user_skills).split(',') if s.strip()]
    missing_skills = []
    skill_counts = {}
    for job_skill in all_job_skills:
        if isinstance(job_skill, str):
            job_skill_lower = job_skill.lower().strip()
            if not any(us in job_skill_lower or job_skill_lower in us for us in user_skills_list):
                if job_skill_lower not in [ms.lower() for ms in missing_skills]:
                    missing_skills.append(job_skill)
                    skill_counts[job_skill] = skill_counts.get(job_skill, 0) + 1
    
    top_skill_gap = max(skill_counts.items(), key=lambda x: x[1])[0] if skill_counts else "Cloud Infrastructure (AWS)"
    
    # Metric 4: Recommended Accreditation
    recommended_accreditation = "PMP or Scrum Master"  # Default, could be enhanced with AI analysis
    
    # Display 4 metrics in columns
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            label="Est. Market Salary Band",
            value=f"HKD {salary_min//1000}k - {salary_max//1000}k / mth",
            delta="+5% vs current",
            delta_color="normal"
        )
    
    with col2:
        st.metric(
            label="Target Role Seniority",
            value=seniority,
            delta="Ready for step up",
            delta_color="off"
        )
    
    with col3:
        st.metric(
            label="Top Skill Gap",
            value=top_skill_gap,
            delta="High Demand in HK",
            delta_color="inverse"
        )
    
    with col4:
        st.metric(
            label="Recommended Accreditation",
            value=recommended_accreditation,
            delta="Unlock 15% more roles",
            delta_color="off"
        )

def display_ranked_matches_table(matched_jobs, user_profile):
    """Display Smart Ranked Matches Table with interactive dataframe"""
    if not matched_jobs:
        return
    
    st.markdown("---")
    st.markdown("### Top AI-Ranked Opportunities")
    
    # Ensure all results have skill match scores calculated
    user_skills = user_profile.get('skills', '')
    
    # Helper function to calculate skill match
    def calc_skill_match(user_skills_str, job_skills_list):
        if not user_skills_str or not job_skills_list:
            return 0.0, []
        user_skills_lower = [s.lower().strip() for s in str(user_skills_str).split(',') if s.strip()]
        job_skills_lower = [s.lower().strip() for s in job_skills_list if isinstance(s, str) and s.strip()]
        if not user_skills_lower or not job_skills_lower:
            return 0.0, []
        matched_skills = []
        for job_skill in job_skills_lower:
            for user_skill in user_skills_lower:
                if job_skill in user_skill or user_skill in job_skill:
                    matched_skills.append(job_skill)
                    break
        match_score = len(matched_skills) / len(job_skills_lower) if job_skills_lower else 0.0
        missing_skills = [s for s in job_skills_lower if s not in matched_skills]
        return min(match_score, 1.0), missing_skills[:5]
    
    for result in matched_jobs:
        if 'skill_match_score' not in result:
            job_skills = result['job'].get('skills', [])
            skill_score, missing_skills = calc_skill_match(user_skills, job_skills)
            result['skill_match_score'] = skill_score
            result['missing_skills'] = missing_skills
    
    # Create DataFrame
    table_data = []
    for i, result in enumerate(matched_jobs):
        job = result['job']
        semantic_score = result['similarity_score']
        skill_score = result.get('skill_match_score', 0.0)
        overall_score = (semantic_score + skill_score) / 2
        
        # Get key matching skills (first 3-4 skills from job that user has)
        job_skills = job.get('skills', [])
        matching_skills = []
        user_skills_list = [s.lower().strip() for s in str(user_skills).split(',') if s.strip()]
        for js in job_skills[:6]:
            if isinstance(js, str):
                js_lower = js.lower().strip()
                if any(us in js_lower or js_lower in us for us in user_skills_list):
                    matching_skills.append(js)
                    if len(matching_skills) >= 4:
                        break
        
        missing_critical = result.get('missing_skills', [])
        missing_critical_skill = missing_critical[0] if missing_critical else "None"
        
        table_data.append({
            'Match Score': int(overall_score * 100),
            'Job Title': job['title'],
            'Company': job['company'],
            'Location': job['location'],
            'Key Matching Skills': matching_skills[:4] if matching_skills else [],
            'Missing Critical Skill': missing_critical_skill,
            '_index': i  # Internal index for selection
        })
    
    df = pd.DataFrame(table_data)
    
    # Configure column display
    column_config = {
        'Match Score': st.column_config.ProgressColumn(
            'AI Match Score',
            help='Overall match percentage',
            min_value=0,
            max_value=100,
            format='%d%%'
        ),
        'Job Title': st.column_config.TextColumn(
            'Job Title',
            width='medium'
        ),
        'Company': st.column_config.TextColumn(
            'Company',
            width='medium'
        ),
        'Location': st.column_config.TextColumn(
            'Location',
            width='small'
        ),
        'Key Matching Skills': st.column_config.ListColumn(
            'Key Matching Skills',
            help='Skills you have that match this role'
        ),
        'Missing Critical Skill': st.column_config.TextColumn(
            'Missing Critical Skill',
            help='Most important skill gap for this role',
            width='medium'
        ),
        '_index': st.column_config.NumberColumn(
            '_index',
            width='small',
            help=None
        )
    }
    
    # Display dataframe with selection
    selected_rows = st.dataframe(
        df,
        column_config=column_config,
        hide_index=True,
        use_container_width=True,
        on_select="rerun",
        selection_mode="single-row"
    )
    
    # Store selected job index
    if selected_rows.selection.rows:
        selected_idx = df.iloc[selected_rows.selection.rows[0]]['_index']
        st.session_state.selected_job_index = int(selected_idx)
    else:
        st.session_state.selected_job_index = None

def display_match_breakdown(matched_jobs, user_profile):
    """Display Match Breakdown & Application Copilot in expander"""
    if st.session_state.selected_job_index is None:
        return
    
    selected_result = matched_jobs[st.session_state.selected_job_index]
    job = selected_result['job']
    semantic_score = selected_result['similarity_score']
    skill_score = selected_result.get('skill_match_score', 0.0)
    missing_skills = selected_result.get('missing_skills', [])
    
    # Calculate skill overlap
    user_skills = user_profile.get('skills', '')
    job_skills = job.get('skills', [])
    user_skills_list = [s.lower().strip() for s in str(user_skills).split(',') if s.strip()]
    job_skills_list = [s.lower().strip() for s in job_skills if isinstance(s, str) and s.strip()]
    
    matched_skills_count = 0
    for js in job_skills_list:
        if any(us in js or js in us for us in user_skills_list):
            matched_skills_count += 1
    
    total_required = len(job_skills_list) if job_skills_list else 1
    skill_overlap_pct = (matched_skills_count / total_required * 100) if total_required > 0 else 0
    
    # Expander title
    expander_title = f"Deep Dive: {job['title']} at {job['company']}"
    
    with st.expander(expander_title, expanded=True):
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("#### Why this is a fit")
            
            # Score breakdown
            st.markdown(f"""
            **Semantic Score:** {semantic_score:.0%}  
            Your experience contextually aligns closely with role requirements.
            
            **Skill Overlap:** {skill_overlap_pct:.0f}%  
            You have {matched_skills_count}/{total_required} required core skills.
            """)
            
            # Recruiter Note
            if semantic_score >= 0.7:
                recruiter_note = f"This role heavily emphasizes recent experience in {job.get('skills', ['relevant skills'])[0] if job.get('skills') else 'relevant skills'}, which is a strong point in your profile."
            else:
                recruiter_note = "Consider highlighting more relevant experience from your background to strengthen your application."
            
            st.info(f"**Recruiter Note:** {recruiter_note}")
        
        with col2:
            st.markdown("#### Application Copilot")
            
            # Accreditation Gap Action
            if missing_skills:
                top_missing = missing_skills[0]
                # Check if it's a certification-related skill
                cert_keywords = ['certification', 'certified', 'accreditation', 'license', 'pmp', 'scrum', 'hkicpa', 'cpa']
                is_cert = any(kw in top_missing.lower() for kw in cert_keywords)
                
                if is_cert:
                    st.warning(f"⚠️ **Crucial Gap:** This job highly values {top_missing}. Consider starting this certification.")
                else:
                    st.warning(f"⚠️ **Skill Gap:** Consider developing expertise in {top_missing}.")
            
            # Primary Action Button
            if st.button("✨ Tailor Resume for this Job", use_container_width=True, type="primary", key="tailor_resume_button"):
                st.session_state.selected_job = job
                st.session_state.show_resume_generator = True
                st.rerun()
            
            st.caption("Generates a citation-locked, AI-optimized CV emphasizing your matching skills.")

def format_resume_as_text(resume_data):
    """Format structured resume JSON as plain text"""
    text = []
    
    # Header
    header = resume_data.get('header', {})
    if header.get('name'):
        text.append(header['name'].upper())
        text.append("")
    
    # Contact info
    contact = []
    if header.get('email'):
        contact.append(header['email'])
    if header.get('phone'):
        contact.append(header['phone'])
    if header.get('location'):
        contact.append(header['location'])
    if header.get('linkedin'):
        contact.append(header['linkedin'])
    if header.get('portfolio'):
        contact.append(header['portfolio'])
    
    if contact:
        text.append(' | '.join(contact))
        text.append("")
    
    # Title
    if header.get('title'):
        text.append(header['title'])
        text.append("")
    
    # Summary
    if resume_data.get('summary'):
        text.append("PROFESSIONAL SUMMARY")
        text.append("-" * 50)
        text.append(resume_data['summary'])
        text.append("")
    
    # Skills
    skills = resume_data.get('skills_highlighted', [])
    if skills:
        text.append("KEY SKILLS")
        text.append("-" * 50)
        text.append(' • '.join(skills))
        text.append("")
    
    # Experience
    experience = resume_data.get('experience', [])
    if experience:
        text.append("PROFESSIONAL EXPERIENCE")
        text.append("-" * 50)
        for exp in experience:
            exp_line = exp.get('title', '')
            if exp.get('company'):
                exp_line += f" at {exp['company']}"
            if exp.get('dates'):
                exp_line += f" | {exp['dates']}"
            text.append(exp_line)
            
            bullets = exp.get('bullets', [])
            for bullet in bullets:
                if bullet.strip():
                    text.append(f"  • {bullet}")
            text.append("")
    
    # Education
    if resume_data.get('education'):
        text.append("EDUCATION")
        text.append("-" * 50)
        text.append(resume_data['education'])
        text.append("")
    
    # Certifications
    if resume_data.get('certifications'):
        text.append("CERTIFICATIONS & AWARDS")
        text.append("-" * 50)
        text.append(resume_data['certifications'])
    
    return '\n'.join(text)

def main():
    # Check if resume generator should be shown
    if st.session_state.get('show_resume_generator', False):
        display_resume_generator()
        return
    
    # Render sidebar with controls
    render_sidebar()
    
    # Main dashboard area - only show after analysis
    if not st.session_state.get('dashboard_ready', False) or not st.session_state.matched_jobs:
        # Show empty state
        st.info("👆 Upload your CV in the sidebar and click 'Analyze Profile & Find Matches' to see your market positioning and ranked opportunities.")
        return
    
    # Display Market Positioning Profile (Top Section)
    display_market_positioning_profile(
        st.session_state.matched_jobs,
        st.session_state.user_profile
    )
    
    # Display Smart Ranked Matches Table (Middle Section)
    display_ranked_matches_table(
        st.session_state.matched_jobs,
        st.session_state.user_profile
    )
    
    # Display Match Breakdown & Application Copilot (Bottom Section)
    display_match_breakdown(
        st.session_state.matched_jobs,
        st.session_state.user_profile
    )

if __name__ == "__main__":
    main()
