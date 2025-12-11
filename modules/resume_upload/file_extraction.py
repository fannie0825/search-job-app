"""File extraction functions for resume upload"""
import streamlit as st
import PyPDF2
from docx import Document


def extract_text_from_resume(uploaded_file):
    """Extract text from uploaded resume file (PDF, DOCX, or TXT)"""
    try:
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        if file_type == 'pdf':
            uploaded_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        
        elif file_type == 'docx':
            uploaded_file.seek(0)
            doc = Document(uploaded_file)
            
            text_parts = []
            
            # Extract text from paragraphs in the main body
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables (very common in resumes for formatting)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            # Extract text from headers (if any)
            for section in doc.sections:
                header = section.header
                if header:
                    for paragraph in header.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
                footer = section.footer
                if footer:
                    for paragraph in footer.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
            
            text = "\n".join(text_parts)
            return text
        
        elif file_type == 'txt':
            uploaded_file.seek(0)
            text = str(uploaded_file.read(), "utf-8")
            return text
        
        else:
            st.error(f"Unsupported file type: {file_type}. Please upload PDF, DOCX, or TXT.")
            return None
            
    except Exception as e:
        st.error(f"Error extracting text from resume: {e}")
        return None
